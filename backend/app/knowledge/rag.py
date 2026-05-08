"""RAG pipeline: document upload → MinIO → chunk → Qdrant."""

import logging
import os
import uuid
from io import BytesIO

from minio import Minio
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.knowledge_document import KnowledgeDocument

logger = logging.getLogger(__name__)

COLLECTION_NAME = "spawnhive_docs"
MEMORY_COLLECTION_NAME = "memory_entities"
CHUNK_SIZE = 500  # characters
CHUNK_OVERLAP = 50
DEFAULT_FASTEMBED_DIM = 384  # BAAI/bge-small-en-v1.5 dimension


def get_minio_client() -> Minio:
    settings = get_settings()
    return Minio(
        settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=False,
    )


def get_qdrant_client() -> QdrantClient:
    settings = get_settings()
    return QdrantClient(url=settings.qdrant_url)


def ensure_bucket(client: Minio, bucket: str = "spawnhive"):
    if not client.bucket_exists(bucket):
        client.make_bucket(bucket)


def ensure_collection(client: QdrantClient, dim: int, name: str = COLLECTION_NAME):
    """Create collection with given dim if absent. Raises if existing dim mismatches."""
    collections = [c.name for c in client.get_collections().collections]
    if name not in collections:
        client.create_collection(
            collection_name=name,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
        )
        return
    info = client.get_collection(name)
    existing_dim = info.config.params.vectors.size
    if existing_dim != dim:
        raise RuntimeError(
            f"Qdrant collection '{name}' has dim={existing_dim} but embedder "
            f"produced dim={dim}. Reset RAG (POST /api/knowledge/reset) to switch providers."
        )


def extract_text(filename: str, content: bytes) -> str:
    """Extract plaintext from a document by extension. Falls back to UTF-8 decode."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        from docx import Document
        doc = Document(BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text)
    return content.decode("utf-8", errors="ignore")


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - CHUNK_OVERLAP
    return chunks


async def get_embeddings(texts: list[str]) -> list[list[float]]:
    """Thin shim: delegate to the configured EmbeddingProvider plugin."""
    from app.plugins.embeddings import get_embedding_provider

    return await get_embedding_provider().embed(texts)


async def process_document(
    db: AsyncSession,
    filename: str,
    content: bytes,
    workspace_id: uuid.UUID,
) -> KnowledgeDocument:
    """Upload to MinIO, chunk, embed, index in Qdrant. Workspace-scoped."""
    doc_id = uuid.uuid4()
    s3_path = f"documents/{workspace_id}/{doc_id}/{filename}"

    # 1. Upload to MinIO
    minio_client = get_minio_client()
    ensure_bucket(minio_client)

    minio_client.put_object(
        "spawnhive", s3_path, BytesIO(content), len(content),
    )
    logger.info(f"Uploaded {filename} to MinIO: {s3_path}")

    # 2. Extract text and chunk
    text = extract_text(filename, content)
    chunks = chunk_text(text)

    if not chunks:
        # Save document record even with no chunks
        doc = KnowledgeDocument(
            id=doc_id, filename=filename, s3_path=s3_path, chunk_count=0,
            workspace_id=workspace_id,
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        return doc

    # 3. Generate embeddings
    embeddings = await get_embeddings(chunks)

    # 4. Index in Qdrant
    qdrant = get_qdrant_client()
    ensure_collection(qdrant, dim=len(embeddings[0]))

    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=emb,
            payload={
                "document_id": str(doc_id),
                "filename": filename,
                "chunk_index": i,
                "text": chunk,
                "workspace_id": str(workspace_id),
            },
        )
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings))
    ]

    qdrant.upsert(collection_name=COLLECTION_NAME, points=points)
    logger.info(f"Indexed {len(points)} chunks for {filename}")

    # 5. Save to DB
    doc = KnowledgeDocument(
        id=doc_id, filename=filename, s3_path=s3_path, chunk_count=len(chunks),
        workspace_id=workspace_id,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    return doc


async def delete_document_data(db: AsyncSession, doc: KnowledgeDocument):
    """Delete document from MinIO, Qdrant, and DB."""
    # MinIO
    try:
        minio_client = get_minio_client()
        minio_client.remove_object("spawnhive", doc.s3_path)
    except Exception as e:
        logger.warning(f"MinIO delete failed: {e}")

    # Qdrant
    try:
        qdrant = get_qdrant_client()
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        qdrant.delete(
            collection_name=COLLECTION_NAME,
            points_selector=Filter(
                must=[FieldCondition(key="document_id", match=MatchValue(value=str(doc.id)))]
            ),
        )
    except Exception as e:
        logger.warning(f"Qdrant delete failed: {e}")

    # DB
    await db.delete(doc)
    await db.commit()


async def reset_collection(db: AsyncSession, workspace_id: uuid.UUID) -> dict:
    """Delete documents and memory entities for a single workspace.

    For Qdrant, we delete only the points whose payload.workspace_id matches.
    Collections themselves remain so other workspaces aren't affected.
    """
    from sqlalchemy import select
    from qdrant_client.models import FieldCondition, Filter, MatchValue
    from app.models.memory import MemoryEntity, MemoryRelation

    ws_filter = Filter(must=[
        FieldCondition(key="workspace_id", match=MatchValue(value=str(workspace_id))),
    ])
    qdrant = get_qdrant_client()
    for col in (COLLECTION_NAME, MEMORY_COLLECTION_NAME):
        try:
            qdrant.delete(collection_name=col, points_selector=ws_filter)
            logger.info(f"Cleared workspace {workspace_id} from Qdrant collection {col}")
        except Exception as e:
            logger.warning(f"Qdrant cleanup of {col} failed (may not exist): {e}")

    minio_client = get_minio_client()
    s3_paths_removed = 0
    try:
        prefix = f"documents/{workspace_id}/"
        for obj in minio_client.list_objects("spawnhive", prefix=prefix, recursive=True):
            minio_client.remove_object("spawnhive", obj.object_name)
            s3_paths_removed += 1
    except Exception as e:
        logger.warning(f"MinIO cleanup failed: {e}")

    docs = (
        await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.workspace_id == workspace_id)
        )
    ).scalars().all()
    db_count = len(docs)
    for d in docs:
        await db.delete(d)

    relations = (
        await db.execute(
            select(MemoryRelation).where(MemoryRelation.workspace_id == workspace_id)
        )
    ).scalars().all()
    for r in relations:
        await db.delete(r)
    entities = (
        await db.execute(
            select(MemoryEntity).where(MemoryEntity.workspace_id == workspace_id)
        )
    ).scalars().all()
    mem_count = len(entities)
    for e in entities:
        await db.delete(e)

    await db.commit()

    return {
        "docs_deleted": db_count,
        "s3_objects_removed": s3_paths_removed,
        "memory_entities_deleted": mem_count,
    }


async def search_documents(
    query: str, workspace_id: uuid.UUID, limit: int = 5
) -> list[dict]:
    """Search for relevant document chunks via vector similarity (workspace-scoped)."""
    try:
        from qdrant_client.models import FieldCondition, Filter, MatchValue

        embeddings = await get_embeddings([query])
        qdrant = get_qdrant_client()
        ensure_collection(qdrant, dim=len(embeddings[0]))

        results = qdrant.query_points(
            collection_name=COLLECTION_NAME,
            query=embeddings[0],
            limit=limit,
            query_filter=Filter(must=[
                FieldCondition(
                    key="workspace_id",
                    match=MatchValue(value=str(workspace_id)),
                ),
            ]),
        )

        return [
            {
                "text": r.payload.get("text", ""),
                "filename": r.payload.get("filename", ""),
                "chunk_index": r.payload.get("chunk_index", 0),
                "score": r.score,
            }
            for r in results.points
        ]
    except Exception as e:
        logger.error(f"Search failed: {e}")
        return []
